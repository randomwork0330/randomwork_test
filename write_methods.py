"""
한글 논문 Methods 파트 docx 생성 스크립트
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── 페이지 여백 설정 (상하좌우 2.5cm)
section = doc.sections[0]
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.left_margin   = Cm(3.0)
section.right_margin  = Cm(3.0)

# ── 기본 폰트 설정 유틸
def set_font(run, name="맑은 고딕", size=11, bold=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18 if level == 1 else 10)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text)
    size = 14 if level == 1 else 12
    set_font(run, size=size, bold=True, color=(0x1F, 0x49, 0x7D) if level == 1 else (0x2E, 0x74, 0xB5))
    return p

def add_body(doc, text, indent=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after   = Pt(4)
    p.paragraph_format.line_spacing  = Pt(20)
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.7)
    run = p.add_run(text)
    set_font(run, size=10.5)
    return p

def add_equation(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text)
    set_font(run, name="Courier New", size=10.5, bold=True, color=(0x1F, 0x49, 0x7D))
    return p

def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after  = Pt(3)
    p.paragraph_format.left_indent  = Cm(1.0)
    run = p.add_run(text)
    set_font(run, size=10.5)
    return p

# ════════════════════════════════════════════════
# 제목
# ════════════════════════════════════════════════
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.paragraph_format.space_after = Pt(4)
r = title.add_run("연구 방법 (Methods)")
set_font(r, size=16, bold=True, color=(0x1F, 0x49, 0x7D))

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.paragraph_format.space_after = Pt(20)
r2 = sub.add_run("이종 생물학적 네트워크 기반 Random Walk with Restart를 활용한 핵심 유전자·약물 노드 중요도 추론")
set_font(r2, size=11, bold=False, color=(0x59, 0x59, 0x59))

doc.add_paragraph()

# ════════════════════════════════════════════════
# 1. 데이터 수집 및 전처리
# ════════════════════════════════════════════════
add_heading(doc, "1. 데이터 수집 및 전처리", level=1)

add_body(doc,
    "본 연구에서는 세 종류의 이종(heterogeneous) 생물학적 데이터를 수집하여 "
    "통합 네트워크 분석에 활용하였다. 각 데이터 소스는 유전자 간 연관성, "
    "단백질 간 물리적 상호작용, 그리고 약물-유전자 표적 관계를 포함하며, "
    "아래와 같이 구성된다.", indent=True)

add_heading(doc, "1.1 유전자-유전자 연관성 데이터 (GGA, _A_)", level=2)
add_body(doc,
    "TMPRSS2-ERG 융합 유전자와의 공발현 관계를 기반으로 한 유전자-유전자 "
    "연관성(gene-gene association, GGA) 데이터를 사용하였다. 각 유전자의 "
    "TMPRSS2-ERG 발현 패턴과의 절대 피어슨 상관계수(absolute Pearson correlation "
    "coefficient, absPearsonR)를 연관성 강도 지표로 활용하였으며, "
    "데이터 형식은 gene | absPearsonR 의 두 열로 구성된다.", indent=True)

add_heading(doc, "1.2 단백질-단백질 상호작용 데이터 (PPI, _B_)", level=2)
add_body(doc,
    "STRING 데이터베이스(v11 이상)로부터 인간(taxon ID: 9606) 단백질 간 물리적 "
    "상호작용(physical PPI) 데이터를 수집하였다. 각 상호작용 쌍(gene1, gene2)에 "
    "대해 STRING이 제공하는 통합 신뢰도 점수(combined score)를 0~1 범위로 "
    "사전 정규화하여 사용하였으며, 데이터 형식은 gene1 | gene2 | score 로 구성된다.", indent=True)

add_heading(doc, "1.3 약물-유전자 표적 데이터 (DGT, _C_)", level=2)
add_body(doc,
    "임상 유전체 데이터베이스 CIViC(Clinical Interpretation of Variants in Cancer)로부터 "
    "약물 감수성(sensitivity) 근거를 기반으로 한 약물-유전자 표적 데이터를 수집하였다. "
    "각 항목은 약물명(drug), 표적 유전자(target_gene), 최종 표적 유전자(final_target_gene), "
    "그리고 근거 점수(score)로 구성되며, final_target_gene 이 존재하는 경우를 우선적으로 "
    "사용하고 그렇지 않은 경우 target_gene 을 대체 사용하였다.", indent=True)

add_heading(doc, "1.4 데이터 필터링", level=2)
add_body(doc,
    "이종 데이터 간 노드 집합의 일관성을 유지하기 위해, GGA 데이터(_A_)에 포함된 "
    "유전자 집합(gene_set)을 기준 집합으로 설정하였다. PPI 데이터는 gene1 또는 "
    "gene2 중 하나 이상이 기준 집합에 포함된 상호작용만을 선택하였으며, "
    "DGT 데이터는 final_target_gene(또는 target_gene)이 기준 집합에 포함된 "
    "항목만을 필터링하여 사용하였다.", indent=True)

# ════════════════════════════════════════════════
# 2. 정규화
# ════════════════════════════════════════════════
add_heading(doc, "2. 엣지 가중치 정규화: Z-score → Sigmoid 변환", level=1)

add_body(doc,
    "세 데이터 소스는 측정 방식과 값의 분포가 서로 상이한 이종 데이터이므로, "
    "단순 Min-Max 정규화를 적용할 경우 극단값(outlier) 하나가 전체 스케일을 왜곡하고 "
    "분포 정보가 손실되는 문제가 발생한다. 이를 극복하기 위해 본 연구에서는 "
    "Z-score 표준화 이후 Sigmoid 함수를 적용하는 2단계 확률적 정규화 방법을 채택하였다.", indent=True)

add_heading(doc, "2.1 1단계: Z-score 표준화", level=2)
add_body(doc,
    "각 데이터 소스 내의 엣지 가중치 x에 대해 해당 소스 내 평균(μ)과 표준편차(σ)를 "
    "이용하여 Z-score를 산출함으로써, 소스 간 분포 편향 및 스케일 차이를 제거한다.")

add_equation(doc, "z = (x − μ) / σ")

add_body(doc,
    "단, σ = 0인 경우(모든 값이 동일) z = 0으로 처리하여 이후 Sigmoid 변환 시 "
    "0.5를 반환하도록 하였다.")

add_heading(doc, "2.2 2단계: Sigmoid 변환", level=2)
add_body(doc,
    "Z-score 변환된 값 z에 표준 로지스틱 시그모이드 함수를 적용하여 "
    "0~1 범위의 확률값 p로 변환한다.")

add_equation(doc, "p = 1 / (1 + exp(−z))")

add_body(doc,
    "이 변환을 통해 평균에 가까운 값은 0.5 부근에서 높은 분별력을 가지며, "
    "극단값은 0 또는 1에 점진적으로 수렴하여 outlier의 영향이 자연스럽게 압축된다. "
    "또한 변환된 값은 로지스틱 확률 분포의 관점에서 해석 가능하므로, "
    "이종 데이터 간의 확률론적으로 공정한 비교가 가능하다.", indent=True)

add_heading(doc, "2.3 중복 엣지 병합", level=2)
add_body(doc,
    "GGA 데이터와 PPI 데이터 사이에 동일한 유전자 쌍이 중복으로 존재하는 경우, "
    "두 소스의 정규화된 가중치를 신뢰도 가중 평균으로 병합하였다. "
    "PPI(STRING 물리적 상호작용)가 GGA(전사체 상관관계) 대비 실험적 근거가 "
    "더 강하다는 점을 반영하여 가중치를 각각 0.4, 0.6으로 설정하였다.")

add_equation(doc, "w_merged = 0.4 × w_GGA + 0.6 × w_PPI")

add_body(doc,
    "DGT 데이터에서 동일한 약물-유전자 쌍이 중복으로 나타나는 경우에는 "
    "정규화된 값이 더 큰 항목을 최종 가중치로 채택하였다.")

# ════════════════════════════════════════════════
# 3. 이종 네트워크 구성
# ════════════════════════════════════════════════
add_heading(doc, "3. 이종 생물학적 네트워크 구성", level=1)

add_body(doc,
    "세 데이터 소스를 통합하여 하나의 무방향 가중 이종 그래프 G = (V, E)를 구성하였다. "
    "노드 집합 V는 유전자 노드(gene)와 약물 노드(drug)로 구성되며, "
    "엣지 집합 E는 각 데이터 소스로부터 유래한 세 종류의 엣지로 이루어진다.", indent=True)

add_body(doc, "네트워크의 노드 및 엣지 구성은 다음과 같다.")

add_bullet(doc, "GGA 엣지: ERG 유전자를 앵커(anchor)로 하여 absPearsonR 기반 "
                "공발현 유전자들을 star 형태로 연결 (weight: Z-score→Sigmoid 변환값)")
add_bullet(doc, "PPI 엣지: STRING 물리적 상호작용 기반 유전자 쌍 연결 "
                "(weight: Z-score→Sigmoid 변환값)")
add_bullet(doc, "DGT 엣지: CIViC 근거 기반 약물-유전자 표적 관계 연결 "
                "(weight: Z-score→Sigmoid 변환값)")
add_bullet(doc, "GGA+PPI 중복 엣지: 신뢰도 가중 평균으로 병합 (GGA=0.4, PPI=0.6)")

add_body(doc,
    "모든 엣지 가중치는 Z-score → Sigmoid 정규화를 통해 동일한 0~1 척도로 통일되었으므로, "
    "이종 소스 간 편향 없이 Random Walk의 이동 확률에 반영될 수 있다.", indent=True)

# ════════════════════════════════════════════════
# 4. RWR
# ════════════════════════════════════════════════
add_heading(doc, "4. Random Walk with Restart (RWR) 시뮬레이션", level=1)

add_body(doc,
    "구성된 이종 네트워크에서 각 노드의 중요도를 추론하기 위해 "
    "Random Walk with Restart(RWR) 알고리즘을 확률적 시뮬레이션 방식으로 구현하였다. "
    "RWR은 네트워크 상에서 임의 보행자(random walker)가 확률적으로 이동하면서 "
    "각 노드를 방문하는 빈도를 측정하는 방법으로, 방문 빈도가 높은 노드일수록 "
    "네트워크 내에서 구조적·기능적으로 중요한 위치를 점한다고 해석한다.", indent=True)

add_heading(doc, "4.1 이동 규칙", level=2)
add_body(doc, "매 스텝 t에서 보행자는 다음의 두 가지 규칙 중 하나를 확률적으로 수행한다.")

add_bullet(doc,
    "Teleport (확률 r = 0.15): 시작 노드(start node)로 즉시 복귀한다. "
    "이를 통해 보행자가 네트워크의 특정 지역에 갇히는 현상(local trapping)을 방지하고, "
    "전역적 탐색이 가능하도록 한다.")
add_bullet(doc,
    "이동 (확률 1 − r = 0.85): 현재 노드의 이웃 노드 중 하나로 이동한다. "
    "이동 확률은 엣지 가중치에 비례하는 local-softmax 방식으로 산출된다.")

add_body(doc, "이웃 노드 j로의 이동 확률은 다음과 같이 정의된다.")
add_equation(doc, "P(j | i) = (w_ij + ε) / Σ_k (w_ik + ε)")
add_body(doc,
    "여기서 w_ij는 노드 i와 j 사이의 정규화된 엣지 가중치이며, "
    "ε = 1×10⁻⁸은 가중치가 0인 엣지에도 최소한의 이동 확률을 보장하기 위한 "
    "수치 안정화 항(numerical stability term)이다.")

add_heading(doc, "4.2 시뮬레이션 설정", level=2)
add_body(doc, "본 연구의 RWR 시뮬레이션은 다음의 하이퍼파라미터를 사용하였다.")

add_bullet(doc, "시작 노드: 네트워크 내 모든 노드를 각각 시작점으로 설정")
add_bullet(doc, "워크 횟수(n_walks): 시작 노드당 1,000회")
add_bullet(doc, "스텝 수(walk_length): 워크당 최대 100 스텝")
add_bullet(doc, "Restart 확률(r): 0.15")
add_bullet(doc, "난수 시드(seed): 42 (재현성 확보)")

add_body(doc,
    "총 시뮬레이션 횟수는 |V| × 1,000 회이며, 각 워크에서 방문한 모든 노드의 "
    "방문 횟수를 누적 집계하였다. 최종 방문 횟수를 전체 방문 횟수의 합으로 나누어 "
    "각 노드의 방문 확률(visit probability)을 산출하였다.", indent=True)

add_heading(doc, "4.3 노드 중요도 정의", level=2)
add_body(doc,
    "RWR 시뮬레이션 종료 후, 각 노드 v의 중요도 점수는 해당 노드의 방문 확률로 정의한다.")

add_equation(doc, "Importance(v) = visit_count(v) / Σ_u visit_count(u)")

add_body(doc,
    "방문 확률이 높은 노드는 네트워크 내에서 다수의 이웃과 강한 연결 관계를 가지며, "
    "여러 경로에서 빈번히 경유되는 허브(hub) 노드임을 의미한다. "
    "유전자 노드의 경우 생물학적 핵심 조절 인자로, "
    "약물 노드의 경우 네트워크 내 핵심 유전자를 표적으로 하는 중요 치료 후보로 해석할 수 있다.", indent=True)

# ════════════════════════════════════════════════
# 5. 결과 분석
# ════════════════════════════════════════════════
add_heading(doc, "5. 결과 분석 및 출력", level=1)

add_body(doc,
    "RWR 시뮬레이션 결과는 노드 타입별(유전자/약물)로 구분하여 분석하였다. "
    "전체 노드 중요도 Top-20을 방문 횟수 기준으로 정렬하여 출력하며, "
    "유전자 노드 Top-10과 약물 노드 Top-10을 각각 별도로 제시하였다. "
    "최종 결과는 노드명, 노드 타입, 방문 횟수, 방문 확률, 연결 차수(degree)를 포함하는 "
    "TSV 파일(node_importance.tsv)로 저장하였다.", indent=True)

# ════════════════════════════════════════════════
# 저장
# ════════════════════════════════════════════════
out_path = "C:/Users/user/Desktop/CNN_CVPR/__DATAAn__/random_work_test/methods_RWR.docx"
doc.save(out_path)
print(f"[Done] 저장 완료: {out_path}")
