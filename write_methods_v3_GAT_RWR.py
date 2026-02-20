"""
method[날짜-시-분]_GAT-RWR.docx 생성 스크립트 (ver3)
GAT+RWR ver3 방법론 논문 메소드 파트 한글 작성
"""

from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

BASE_PATH = r"C:/Users/user/Desktop/CNN_CVPR/__DATAAn__/random_work_test/"

now      = datetime.now()
filename = now.strftime("method%Y%m%d-%H-%M_GAT-RWR.docx")
out_path = os.path.join(BASE_PATH, filename)

doc = Document()

# ── 스타일 헬퍼 ─────────────────────────────────────────────────────────────
def heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    return p

def body(doc, text, bold_kw=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if bold_kw:
        parts = text.split(bold_kw)
        for i, part in enumerate(parts):
            run = p.add_run(part)
            run.font.size = Pt(10.5)
            if i < len(parts) - 1:
                br = p.add_run(bold_kw)
                br.bold = True
                br.font.size = Pt(10.5)
    else:
        run = p.add_run(text)
        run.font.size = Pt(10.5)
    return p

def add_table(doc, headers, rows):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Table Grid'
    hrow = t.rows[0]
    for i, h in enumerate(headers):
        hrow.cells[i].text = h
        hrow.cells[i].paragraphs[0].runs[0].bold = True
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            t.rows[ri+1].cells[ci].text = str(val)
    return t

# ════════════════════════════════════════════════════════════════════════════
# 제목
# ════════════════════════════════════════════════════════════════════════════
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run(
    "이종 생물학적 네트워크 기반 Graph Attention Network + Random Walk with Restart를 이용한\n"
    "전립선암 핵심 유전자 중요도 추론 방법론 (ver3)"
)
r.bold = True
r.font.size = Pt(14)
r.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.add_run(f"[업데이트: {now.strftime('%Y년 %m월 %d일 %H시 %M분')}]").font.size = Pt(10)

doc.add_paragraph()

# ════════════════════════════════════════════════════════════════════════════
# 1. 개요
# ════════════════════════════════════════════════════════════════════════════
heading(doc, "1. 연구 개요 및 ver3 업데이트 사항")
body(doc,
    "본 연구는 TMPRSS2-ERG 융합 전립선암(ergPRAD)의 분자 메커니즘을 이종 생물학적 네트워크 분석으로 "
    "규명하기 위해, Graph Attention Network (GAT)와 Random Walk with Restart (RWR)를 결합한 "
    "통합 계산 분석 파이프라인을 구축하였다. ver3에서는 유전자-유전자 연관 데이터(GGA)를 "
    "ERG 중심의 단방향 상관 벡터에서 TMPRSS2-ERG PRAD TCGA 발현 기반 전체 유전자 쌍(pairwise) "
    "Pearson 상관계수 행렬로 전면 교체하여, 유전자 간 공동발현 관계를 보다 포괄적으로 반영하였다."
)

doc.add_paragraph()
body(doc,
    "ver2 대비 주요 변경 사항은 다음과 같다. ① GGA 데이터: ERG vs. 개별 유전자 1:N 방식에서 "
    "3,871개 유전자 간 모든 쌍(3,267,720 쌍) absPearsonR 기반 N:N pairwise 방식으로 교체. "
    "② PPI(_B_), DGT(_C_) 데이터 및 GAT 아키텍처, RWR 파라미터는 ver2와 동일하게 유지. "
    "③ 결과 파일: node_importance_v3.tsv 신규 생성."
)

# ════════════════════════════════════════════════════════════════════════════
# 2. 데이터
# ════════════════════════════════════════════════════════════════════════════
heading(doc, "2. 입력 데이터")

heading(doc, "2.1 유전자-유전자 공동발현 연관 네트워크 (_A_, 신규)", level=2)
body(doc,
    "TCGA PRAD 코호트 중 TMPRSS2-ERG 양성 50개 및 음성 50개 샘플, 총 100개 샘플의 RNA-seq 발현 "
    "데이터(19,928개 유전자 × 100 샘플)를 사용하였다. 분석은 2단계로 수행하였다."
)
body(doc,
    "1단계: ERG와 Pearson 절대 상관계수(absPearsonR) ≥ 0.3 인 유전자를 추출하여 ERG 공동발현 "
    "유전자 집합(3,871개)을 선별하였다. ERG를 기준 벡터로 사용하고, 전체 유전자 행렬을 "
    "행(유전자)별 Z-score 표준화한 후 행렬 내적으로 일괄 계산하였다."
)
body(doc,
    "2단계: 선별된 3,871개 유전자 간 모든 쌍에 대해 Pearson 상관계수를 행렬 연산으로 계산하고 "
    "(3,871 × 3,871 행렬), absPearsonR ≥ 0.3 인 3,267,720 쌍을 최종 GGA 데이터셋으로 확정하였다. "
    "각 유전자 쌍의 absPearsonR을 에지 가중치의 원시 값으로 사용하였다."
)

add_table(doc,
    ["항목", "값"],
    [
        ["분석 코호트", "TCGA PRAD (ERG양성 50 + ERG음성 50, 총 100샘플)"],
        ["입력 유전자 수", "19,928개"],
        ["Step1 ERG 상관 유전자", "3,871개 (absPearsonR ≥ 0.3)"],
        ["Step2 최종 GGA 쌍", "3,267,720쌍 (absPearsonR ≥ 0.3)"],
        ["absPearsonR 범위", "0.300 ~ 0.984 (중앙값 0.436)"],
        ["파일", "__A__gene_gene_asso.tsv (gene_A, gene_B, pearsonR, absPearsonR)"],
    ]
)

heading(doc, "2.2 단백질-단백질 상호작용 (_B_, ver2 동일)", level=2)
body(doc,
    "STRING DB v12 Human (9606) Physical interaction 데이터를 사용하였다. "
    "전체 1,477,610개 상호작용 중 GGA 유전자 집합에 속하는 유전자를 포함하는 "
    "539,804개 쌍을 필터링하였다. STRING combined score를 에지 가중치로 사용하였다."
)

heading(doc, "2.3 약물-유전자 표적 관계 (_C_, ver2 동일)", level=2)
body(doc,
    "CIViC(Clinical Interpretation of Variants in Cancer) 데이터베이스에서 "
    "임상적으로 검증된 drug-gene target 관계 1,628개를 수집하였으며, "
    "GGA 유전자 집합에 속하는 유전자를 표적으로 하는 322개 약물-유전자 쌍을 사용하였다."
)

# ════════════════════════════════════════════════════════════════════════════
# 3. 이종 네트워크 구성
# ════════════════════════════════════════════════════════════════════════════
heading(doc, "3. 이종 생물학적 네트워크 구성")
body(doc,
    "GGA, PPI, DGT 데이터를 통합하여 무방향 가중 이종 그래프(undirected weighted heterogeneous "
    "graph) G = (V, E)를 구성하였다. 노드 집합 V는 유전자 노드와 약물 노드로 구성되며, "
    "에지 집합 E는 세 가지 에지 유형(GGA, PPI, DGT)을 포함한다."
)

heading(doc, "3.1 에지 가중치 정규화 (Z-score → Sigmoid)", level=2)
body(doc,
    "각 에지 유형별로 독립적인 Z-score → Sigmoid 정규화를 적용하여 에지 가중치를 0~1 범위의 "
    "확률값으로 변환하였다. 이는 서로 다른 스케일을 갖는 이종 데이터를 통합 비교 가능하게 "
    "하기 위함이다."
)
body(doc, "정규화 수식: z = (x − μ) / σ,  p = 1 / (1 + exp(−z))")
body(doc, "여기서 μ, σ는 각 에지 유형 내의 평균과 표준편차이다.")

heading(doc, "3.2 중복 에지 병합 (GGA + PPI)", level=2)
body(doc,
    "동일한 유전자 쌍이 GGA와 PPI 양쪽에 존재할 경우, 가중 평균으로 병합하였다: "
    "w_merged = 0.4 × w_GGA + 0.6 × w_PPI. "
    "PPI 가중치에 더 높은 가중치를 부여하는 이유는 PPI가 실험적으로 검증된 직접적인 "
    "단백질 상호작용을 나타내기 때문이다."
)

heading(doc, "3.3 네트워크 규모 (ver3)", level=2)
add_table(doc,
    ["에지 유형", "에지 수", "비고"],
    [
        ["GGA (gene-gene pairwise)", "~3,267,720", "absPearsonR Z-sig 정규화"],
        ["PPI (gene-gene)", "~539,804", "STRING physical score"],
        ["DGT (drug-gene)", "~322", "CIViC 임상 검증"],
        ["GGA+PPI 중복 병합", "해당 쌍 수", "weighted avg (0.4:0.6)"],
        ["총 노드", "~4,058", "유전자 ~3,871 + 약물 ~187"],
        ["총 에지", "~3,807,846", "중복 병합 후 단방향 기준"],
    ]
)

# ════════════════════════════════════════════════════════════════════════════
# 4. 노드 특징 행렬
# ════════════════════════════════════════════════════════════════════════════
heading(doc, "4. 노드 특징 행렬 (ver2 동일)")
body(doc,
    "각 노드 v에 대해 8차원 특징 벡터 x_v ∈ R^8을 구성하였다."
)
add_table(doc,
    ["차원", "특징", "설명"],
    [
        ["0", "is_gene", "유전자 노드 여부 (0/1)"],
        ["1", "is_drug", "약물 노드 여부 (0/1)"],
        ["2", "log_degree", "log(1 + degree)"],
        ["3", "gga_ratio", "GGA 에지 비율"],
        ["4", "ppi_ratio", "PPI 에지 비율"],
        ["5", "dgt_ratio", "DGT 에지 비율"],
        ["6", "mean_weight", "인접 에지 평균 가중치"],
        ["7", "max_weight", "인접 에지 최대 가중치"],
    ]
)

# ════════════════════════════════════════════════════════════════════════════
# 5. GAT 아키텍처
# ════════════════════════════════════════════════════════════════════════════
heading(doc, "5. Graph Attention Network (GAT) 아키텍처 (ver2 동일)")
body(doc,
    "본 연구의 GAT는 에지 특징을 Attention 계산에 통합하는 Edge-Aware GAT를 "
    "Pure PyTorch로 구현하였다. 2개의 GAT 레이어와 노드 유형별 별도 입력 투영(type-specific "
    "projection)으로 구성된다."
)

heading(doc, "5.1 Attention 메커니즘", level=2)
body(doc, "에지 (i → j)의 Attention score는 다음과 같이 계산된다:")
body(doc, "e_ij = LeakyReLU( a^T · [W·h_i || W·h_j || W_e·edge_feat_ij] )")
body(doc, "α_ij = softmax_{j ∈ N(i)}( e_ij / √degree(i) )")
body(doc,
    "여기서 W ∈ R^{F_in×F_out}은 노드 변환 행렬, W_e ∈ R^{5×5}는 에지 특징 변환 행렬, "
    "a ∈ R^{2F_out+5}는 Attention 벡터이다. √degree(i)로 나누는 degree-aware scaling을 적용하여 "
    "고차수(hub) 노드의 Attention 희석을 방지하였다."
)

heading(doc, "5.2 모델 구성", level=2)
add_table(doc,
    ["레이어", "입력 차원", "출력 차원", "Head 수", "Aggregation"],
    [
        ["입력 투영", "8", "8", "-", "Type-specific (gene/drug)"],
        ["GAT Layer 1", "8", "16×8=128", "8", "Concat"],
        ["Dropout", "128", "128", "-", "p=0.3"],
        ["GAT Layer 2", "128", "16", "8", "Mean"],
    ]
)

heading(doc, "5.3 학습 전략", level=2)
body(doc,
    "자기지도 학습(self-supervised learning) 방식의 Link Prediction을 목적 함수로 사용하였다. "
    "에지를 Train:Val:Test = 8:1:1로 분할하고, Negative Sampling(1:1 비율)으로 학습하였다. "
    "Epoch당 1회 전체 그래프 Forward Pass 후 배치 크기(4,096) 만큼의 에지를 샘플링하여 Loss를 계산하는 "
    "방식을 적용, ver2 대비 학습 효율을 33배 개선하였다."
)
body(doc,
    "손실 함수: L = −E[(w_pos · log σ(z_u^T z_v))] − E[(w_neg · log σ(−z_u^T z_v))] + λ·||Z||²"
)
body(doc,
    "약물 노드 관련 에지는 가중치 ×10을 적용하여 소수의 약물 노드 학습을 보완하였다. "
    "Gradient clipping(max_norm=1.0)과 Adam optimizer(lr=0.001), "
    "Val loss 기준 Early Stopping(patience=20)을 적용하였다."
)

# ════════════════════════════════════════════════════════════════════════════
# 6. RWR
# ════════════════════════════════════════════════════════════════════════════
heading(doc, "6. GAT Attention 기반 Random Walk with Restart (RWR)")

heading(doc, "6.1 전이 확률 구성", level=2)
body(doc,
    "학습된 GAT의 Layer 2 Attention(α_ij, 8-head 평균)을 각 노드의 이웃 전이 확률로 변환하였다. "
    "GAT Attention과 원본 에지 가중치를 혼합하여 최종 전이 확률을 계산하였다:"
)
body(doc, "p_combined(i→j) = 0.7 × α_ij + 0.3 × w_orig(i,j)")
body(doc,
    "여기서 α_ij는 GAT Attention (src 기준 softmax 정규화), "
    "w_orig는 Z-score→Sigmoid 정규화된 원본 에지 가중치(src 기준 합=1 정규화)이다. "
    "혼합 비율 α=0.7은 GAT 학습 구조를 우선적으로 반영하면서 원본 가중치의 안정성을 보완한다."
)

heading(doc, "6.2 RWR 파라미터 및 시뮬레이션", level=2)
add_table(doc,
    ["파라미터", "값", "설명"],
    [
        ["n_walks", "1,000", "노드당 시뮬레이션 횟수"],
        ["walk_length", "100", "각 walk 당 최대 스텝"],
        ["restart_prob (r)", "0.15", "출발 노드로 재시작 확률"],
        ["n_workers", "12", "멀티프로세싱 병렬 워커 수"],
        ["seed", "42", "재현성"],
        ["혼합 비율 (α)", "0.7", "GAT:원본 = 7:3"],
    ]
)
body(doc,
    "각 시뮬레이션 스텝에서 현재 노드를 방문 카운터에 기록한 후, "
    "restart_prob(0.15) 확률로 출발 노드로 복귀하거나, "
    "그렇지 않으면 전이 확률 p_combined에 따라 이웃 노드로 이동한다. "
    "전체 walk 수는 노드 수 × n_walks이며, 멀티프로세싱(12 workers)으로 병렬 처리한다."
)

heading(doc, "6.3 노드 중요도 정의", level=2)
body(doc,
    "모든 시뮬레이션 완료 후, 각 노드의 누적 방문 횟수(visit_count)를 전체 방문 횟수로 나누어 "
    "방문 확률(visit_prob)을 계산하며, 이를 최종 노드 중요도 지표로 사용하였다:"
)
body(doc, "importance(v) = visit_count(v) / Σ_u visit_count(u)")

# ════════════════════════════════════════════════════════════════════════════
# 7. GAT Hub Score
# ════════════════════════════════════════════════════════════════════════════
heading(doc, "7. GAT Hub Score")
body(doc,
    "네트워크에서 구조적 허브 역할을 정량화하기 위해, 각 노드로 향하는(incoming) "
    "GAT Attention 값의 평균을 gat_hub_score로 정의하였다:"
)
body(doc, "gat_hub_score(v) = (1/|in(v)|) · Σ_{u→v} α_{uv}")
body(doc,
    "gat_hub_score가 높은 노드는 여러 이웃 노드들이 강하게 주목하는 네트워크 허브를 의미하며, "
    "visit_prob와 함께 노드 중요도를 다각도로 평가하는 보조 지표로 활용된다."
)

# ════════════════════════════════════════════════════════════════════════════
# 8. 통계적 타당성 검증
# ════════════════════════════════════════════════════════════════════════════
heading(doc, "8. 통계적 타당성 검증 (ver2 결과 기반)")
body(doc,
    "GAT+RWR 랭킹의 생물학적 타당성을 검증하기 위해, 문헌 기반으로 주석 처리된 "
    "전립선암 표적치료제 표적 유전자(139개, 이 중 ergPRAD 특화 17개)의 랭킹 분포를 "
    "배경 유전자(17,862개)와 6가지 통계 검정으로 비교하였다."
)

add_table(doc,
    ["검정", "ALL PRAD (139개)", "ergPRAD 특화 (17개)"],
    [
        ["Mann-Whitney U", "p = 3.93×10⁻³⁸", "p = 2.91×10⁻¹⁰"],
        ["KS test (D-stat)", "D = 0.521", "D = 0.815"],
        ["Fisher top 1%", "fold = 16.6×, p = 1.10×10⁻²¹", "fold = 41.2×, p = 1.65×10⁻¹⁰"],
        ["Fisher top 10%", "fold = 4.4×, p = 2.72×10⁻²⁵", "fold = 7.6×, p = 1.56×10⁻¹⁰"],
        ["Permutation (10,000회)", "Z = -12.94, p < 10⁻⁵", "Z = -6.10, p < 10⁻⁵"],
        ["GSEA NES", "1.760 (p=0.000)", "1.903 (p=0.000)"],
    ]
)
body(doc,
    "24/24 검정에서 모두 유의한 결과(p < 0.05)를 보였으며, 특히 ergPRAD 특화 17개 유전자는 "
    "상위 20% 이내에 17/17(100%)이 포함되어(fold=5.0×), GAT+RWR이 생물학적으로 의미있는 "
    "노드 중요도를 추론함을 통계적으로 입증하였다. 해당 결과는 ver2 분석에 기반하며, "
    "ver3 완료 후 동일 검증 절차를 적용할 예정이다."
)

# ════════════════════════════════════════════════════════════════════════════
# 9. 전립선암 표적 유전자 주석
# ════════════════════════════════════════════════════════════════════════════
heading(doc, "9. 전립선암 표적치료제 표적 유전자 주석 체계")
body(doc,
    "RWR 랭킹 결과에 임상 및 분자 생물학적 맥락을 부여하기 위해, "
    "문헌 및 CIViC/COSMIC/OncoKB 기반으로 전립선암 표적치료제 표적 유전자를 수동 큐레이션하였다. "
    "총 139개 유전자에 주석을 부여하였으며, TMPRSS2-ERG 융합 전립선암(ergPRAD)에 특이적인 "
    "17개 유전자는 '(ergPRAD)' 접두사로 별도 표기하였다."
)

add_table(doc,
    ["주석 분류", "유전자 수", "대표 유전자", "주요 약물"],
    [
        ["(ergPRAD) 특화", "17", "ERG, EZH2, BRD4, KDM1A, PTEN, MYC", "BET inhibitor, EZH2i, LSD1i"],
        ["AR axis", "8", "AR, FOXA1, SRD5A1/2, CYP17A1", "Enzalutamide, Abiraterone"],
        ["DNA repair / PARP", "14", "BRCA1/2, ATM, CDK12, PALB2, RAD51", "Olaparib, Rucaparib"],
        ["PI3K/AKT/mTOR", "10", "PTEN, PIK3CA/CB, AKT1, MTOR", "Ipatasertib, Everolimus"],
        ["RTK / RAS / MAPK", "14", "ERBB2, EGFR, MET, KRAS, BRAF", "Cabozantinib, T-DXd"],
        ["Cell cycle / CDK", "9", "RB1, CDK4/6, CCND1, AURKA/B", "Palbociclib, Alisertib"],
        ["Epigenetics", "8", "EZH2, BRD2/3/4, KDM1A, HDAC2", "Tazemetostat, JQ1"],
        ["WNT / JAK-STAT / 기타", "18", "CTNNB1, STAT3, TP53, MYC, VEGFA", "다양"],
    ]
)

# ════════════════════════════════════════════════════════════════════════════
# 10. 버전 비교
# ════════════════════════════════════════════════════════════════════════════
heading(doc, "10. 버전별 분석 비교")
add_table(doc,
    ["항목", "ver1 (RWR only)", "ver2 (GAT+RWR)", "ver3 (GAT+RWR, pairwise GGA)"],
    [
        ["GGA 방식", "ERG vs N (단방향)", "ERG vs N (단방향)", "gene-gene pairwise (N×N)"],
        ["GGA 입력", "4,233 유전자", "4,233 유전자", "3,267,720 쌍 (3,871 유전자)"],
        ["GAT", "없음", "Edge-Aware 2L-8H", "Edge-Aware 2L-8H (동일)"],
        ["전이 확률", "원본 weight", "GAT×0.7 + 원본×0.3", "GAT×0.7 + 원본×0.3 (동일)"],
        ["총 노드", "18,211", "18,211", "~4,058"],
        ["총 에지", "~1.5M", "~1.5M", "~3.8M"],
        ["ERBB2 rank", "13위", "12위", "미완료"],
        ["결과 파일", "node_importance.tsv", "node_importance_v2.tsv", "node_importance_v3.tsv"],
    ]
)

# ════════════════════════════════════════════════════════════════════════════
# 11. 구현 환경
# ════════════════════════════════════════════════════════════════════════════
heading(doc, "11. 구현 환경")
add_table(doc,
    ["항목", "사양"],
    [
        ["언어", "Python 3.x"],
        ["주요 라이브러리", "PyTorch, NetworkX, NumPy, Pandas, SciPy"],
        ["GAT 구현", "Pure PyTorch (torch_geometric 미사용)"],
        ["병렬처리", "multiprocessing.Pool, 12 workers"],
        ["운영체제", "Windows (cp949 인코딩)"],
        ["GitHub", "https://github.com/randomwork0330/randomwork_test"],
    ]
)

# ════════════════════════════════════════════════════════════════════════════
# 저장
# ════════════════════════════════════════════════════════════════════════════
doc.save(out_path)
print(f"저장 완료: {out_path}")
