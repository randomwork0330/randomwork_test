"""
한글 논문 Methods 파트 docx 생성 스크립트 - GAT+RWR ver2
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

doc = Document()

# ── 페이지 여백 설정 (상하좌우 2.5cm)
section = doc.sections[0]
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.left_margin   = Cm(3.0)
section.right_margin  = Cm(3.0)

# ── 유틸 함수
def set_font(run, name="맑은 고딕", size=11, bold=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18 if level == 1 else 10)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text)
    size = 14 if level == 1 else 12
    color = (0x1F, 0x49, 0x7D) if level == 1 else (0x2E, 0x74, 0xB5)
    set_font(run, size=size, bold=True, color=color)
    return p

def add_body(doc, text, indent=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.line_spacing = Pt(20)
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.7)
    run = p.add_run(text)
    set_font(run, size=10.5)
    return p

def add_equation(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text)
    set_font(run, name="Courier New", size=10.5, bold=True, color=(0x1F, 0x49, 0x7D))
    return p

def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(1.0)
    run = p.add_run(text)
    set_font(run, size=10.5)
    return p

# ════════════════════════════════════════════════
# 제목
# ════════════════════════════════════════════════
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.paragraph_format.space_after = Pt(4)
r = title.add_run("연구 방법 (Methods)")
set_font(r, size=16, bold=True, color=(0x1F, 0x49, 0x7D))

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.paragraph_format.space_after = Pt(20)
r2 = sub.add_run(
    "Graph Attention Network와 Random Walk with Restart를 결합한 "
    "이종 생물학적 네트워크 기반 핵심 노드 중요도 추론 (GAT-RWR)"
)
set_font(r2, size=11, bold=False, color=(0x59, 0x59, 0x59))
doc.add_paragraph()

# ════════════════════════════════════════════════
# 1. 연구 개요
# ════════════════════════════════════════════════
add_heading(doc, "1. 연구 개요", level=1)
add_body(doc,
    "본 연구에서는 이종(heterogeneous) 생물학적 네트워크 위에서 Graph Attention Network(GAT)와 "
    "Random Walk with Restart(RWR)를 결합한 GAT-RWR 프레임워크를 제안한다. "
    "기존의 단순 엣지 가중치 기반 RWR은 네트워크의 전역적 구조 맥락(global context)을 "
    "반영하지 못하는 한계가 있다. 이를 극복하기 위해 GAT로 학습된 attention weight를 "
    "RWR의 전이 확률에 통합함으로써, 구조적 중요도를 반영한 방문 확률 기반 노드 중요도를 "
    "추론하였다.", indent=True)

add_body(doc,
    "제안 방법의 전체 파이프라인은 다음과 같이 구성된다: "
    "(1) 이종 생물학적 데이터 로드 및 가중 네트워크 구성, "
    "(2) 노드 특징 행렬 및 엣지 특징 행렬 구성, "
    "(3) Edge-Aware GAT 자기지도 학습(Link Prediction), "
    "(4) 학습된 attention weight를 이용한 전이 확률 캐시 구성, "
    "(5) 멀티프로세싱 RWR 시뮬레이션, "
    "(6) 노드 중요도 산출 및 결과 저장.", indent=True)

# ════════════════════════════════════════════════
# 2. 데이터 수집 및 전처리
# ════════════════════════════════════════════════
add_heading(doc, "2. 데이터 수집 및 전처리", level=1)

add_body(doc,
    "본 연구에서는 세 종류의 이종 생물학적 데이터를 통합하여 활용하였다. "
    "모든 데이터는 TSV 형식으로 저장되어 있으며, GGA(_A_), PPI(_B_), "
    "DGT(_C_)의 세 파일로 구성된다.", indent=True)

add_heading(doc, "2.1 유전자-유전자 연관성 (GGA)", level=2)
add_body(doc,
    "TMPRSS2-ERG 융합 유전자와의 절대 피어슨 상관계수(absPearsonR)를 기반으로 한 "
    "유전자-유전자 연관성(gene-gene association, GGA) 데이터를 사용하였다. "
    "GGA 데이터에 포함된 유전자 집합(gene_set)을 전체 분석의 기준 집합으로 설정하였으며, "
    "ERG를 앵커 노드로 하여 star 구조로 연결된다.")

add_heading(doc, "2.2 단백질-단백질 상호작용 (PPI)", level=2)
add_body(doc,
    "STRING 데이터베이스(v11 이상)로부터 수집된 인간 단백질 간 상호작용 데이터를 사용하였다. "
    "gene1 또는 gene2 중 하나 이상이 GGA 기준 집합에 포함된 상호작용만을 필터링하였다.")

add_heading(doc, "2.3 약물-유전자 표적 (DGT)", level=2)
add_body(doc,
    "CIViC 데이터베이스로부터 수집된 약물 감수성 근거 기반 약물-유전자 표적 데이터를 "
    "사용하였다. final_target_gene이 존재하는 경우 이를 우선 사용하고, "
    "그렇지 않은 경우 target_gene을 대체 사용하였다.")

add_heading(doc, "2.4 엣지 가중치 정규화 (Z-score → Sigmoid)", level=2)
add_body(doc,
    "이종 데이터 소스 간 측정 척도 차이를 제거하기 위해, 각 데이터 소스 내부에서 독립적으로 "
    "Z-score 표준화 후 Sigmoid 변환을 적용하는 2단계 정규화를 수행하였다.")
add_equation(doc, "z = (x - mu) / sigma,    p = 1 / (1 + exp(-z))")
add_body(doc,
    "이를 통해 각 소스 내의 분포 정보를 유지하면서 0~1 범위의 확률값으로 변환하였다. "
    "GGA와 PPI 데이터에서 동일한 유전자 쌍이 중복으로 나타나는 경우, "
    "신뢰도 가중 평균(GGA=0.4, PPI=0.6)으로 병합하였다.")
add_equation(doc, "w_merged = 0.4 * w_GGA  +  0.6 * w_PPI")

# ════════════════════════════════════════════════
# 3. 이종 네트워크 구성
# ════════════════════════════════════════════════
add_heading(doc, "3. 이종 생물학적 네트워크 구성", level=1)

add_body(doc,
    "세 데이터 소스를 통합하여 무방향 가중 이종 그래프 G = (V, E)를 구성하였다. "
    "노드 집합 V는 유전자 노드(18,001개)와 약물 노드(210개)로 구성되며(|V| = 18,211), "
    "엣지 집합 E는 GGA, PPI, DGT 엣지를 포함한다(|E| = 288,336).", indent=True)

add_body(doc, "각 엣지에는 다음의 속성이 저장된다.")
add_bullet(doc, "weight: Z-score → Sigmoid 정규화된 엣지 가중치")
add_bullet(doc, "edge_type: 'gga', 'ppi', 'gga+ppi', 'dgt' 중 하나")

# ════════════════════════════════════════════════
# 4. 노드 및 엣지 특징 행렬 구성
# ════════════════════════════════════════════════
add_heading(doc, "4. 노드 및 엣지 특징 행렬 구성", level=1)

add_body(doc,
    "GAT 학습을 위해 그래프를 PyTorch Tensor 형태로 변환하였다. "
    "노드 특징 행렬 X ∈ R^(N × 8)과 "
    "엣지 특징 행렬 E_feat ∈ R^(M × 5)를 구성하였다.", indent=True)

add_heading(doc, "4.1 노드 특징 행렬 X (N × 8)", level=2)
add_body(doc, "각 노드 v에 대해 다음 8가지 특징을 추출하여 행렬을 구성하였다.")
add_bullet(doc, "x_0: node_type_gene — 유전자 노드이면 1.0, 아니면 0.0")
add_bullet(doc, "x_1: node_type_drug — 약물 노드이면 1.0, 아니면 0.0")
add_bullet(doc, "x_2: log_degree — log(1 + degree(v))")
add_bullet(doc, "x_3: gga_edge_ratio — GGA 엣지 수 / 전체 엣지 수")
add_bullet(doc, "x_4: ppi_edge_ratio — PPI 엣지 수 / 전체 엣지 수")
add_bullet(doc, "x_5: dgt_edge_ratio — DGT 엣지 수 / 전체 엣지 수")
add_bullet(doc, "x_6: mean_edge_weight — 이웃 엣지 가중치 평균")
add_bullet(doc, "x_7: max_edge_weight — 이웃 엣지 가중치 최대값")

add_heading(doc, "4.2 엣지 특징 행렬 E_feat (M × 5)", level=2)
add_body(doc, "각 엣지 e = (i, j)에 대해 다음 5가지 특징을 추출하였다.")
add_bullet(doc, "e_0~e_3: 엣지 타입 원-핫 인코딩 [gga, ppi, gga+ppi, dgt]")
add_bullet(doc, "e_4: 정규화된 엣지 가중치 (Z-score → Sigmoid 변환값)")
add_body(doc,
    "엣지 특징은 양방향(i→j, j→i)으로 복제하여 유향 그래프에서의 "
    "방향별 attention 학습이 가능하도록 하였다.")

# ════════════════════════════════════════════════
# 5. GAT 아키텍처
# ════════════════════════════════════════════════
add_heading(doc, "5. Edge-Aware Graph Attention Network (GAT) 아키텍처", level=1)

add_body(doc,
    "본 연구에서는 엣지 특징을 직접 attention 계산에 통합하는 "
    "Edge-Aware GAT를 설계하였다. PyTorch Geometric이나 DGL과 같은 "
    "외부 GNN 라이브러리 없이 순수 PyTorch로 구현하여 환경 의존성을 최소화하였다.", indent=True)

add_heading(doc, "5.1 전체 구조", level=2)
add_body(doc, "2개의 EdgeAwareGATLayer를 순차적으로 적용하는 구조를 채택하였다.")
add_bullet(doc, "Layer 1: F_in=8, F_out=16, heads=8, concat=True → 출력: (N, 128)")
add_bullet(doc, "Layer 2: F_in=128, F_out=16, heads=8, concat=False → 출력: (N, 16)")
add_body(doc,
    "두 레이어 사이에 ELU 활성화 함수와 Dropout(p=0.3)을 적용하였다. "
    "노드 타입(gene/drug)별 별도의 선형 변환(type-specific projection)을 적용하여 "
    "이종 노드 타입의 특징 공간 차이를 명시적으로 처리하였다.")

add_heading(doc, "5.2 Attention 메커니즘", level=2)
add_body(doc,
    "EdgeAwareGATLayer의 attention 계산은 소스 노드 특징, 목적 노드 특징, "
    "엣지 특징을 함께 고려하는 다음의 수식을 따른다.")

add_equation(doc,
    "e_ij = LeakyReLU(a^T · [W·h_i || W·h_j || W_e·edge_feat_ij])")
add_equation(doc,
    "alpha_ij = exp(e_ij) / SUM_k( exp(e_ik) )")
add_equation(doc,
    "h_i' = ELU( SUM_j( alpha_ij · W · h_j ) )")

add_body(doc,
    "여기서 W ∈ R^(F_out × F_in)은 노드 변환 행렬, "
    "W_e ∈ R^(F_out × 5)는 엣지 특징 투영 행렬, "
    "a ∈ R^(3 × F_out)는 attention 벡터이다. "
    "softmax는 소스 노드 i의 이웃 집합 N(i)에 대해 국소적으로 적용된다(local softmax).", indent=True)

add_heading(doc, "5.3 허브 노드 Attention 보정", level=2)
add_body(doc,
    "ERG(degree=4,279)와 같이 연결 차수가 극단적으로 높은 허브(hub) 노드에서는 "
    "attention이 다수의 이웃에 희석(diluted)되어 각 이웃과의 변별력이 저하되는 문제가 발생한다. "
    "이를 방지하기 위해 attention 계산 전 degree-aware scaling을 적용하였다.")
add_equation(doc, "e_ij  <-  e_ij × (1 / sqrt(degree(i)))")
add_body(doc,
    "이 보정을 통해 허브 노드의 attention이 인위적으로 희석되지 않고 "
    "각 이웃과의 실제 연관성 차이를 더 잘 반영할 수 있도록 하였다.")

# ════════════════════════════════════════════════
# 6. 자기지도 학습: Link Prediction
# ════════════════════════════════════════════════
add_heading(doc, "6. 자기지도 학습: Link Prediction", level=1)

add_body(doc,
    "별도의 레이블이 없는 생물학적 네트워크에서 GAT를 학습하기 위해 "
    "자기지도(self-supervised) 방식의 Link Prediction 손실 함수를 사용하였다. "
    "실제로 연결된 엣지(positive sample)와 연결되지 않은 임의의 쌍(negative sample)을 "
    "구분하는 이진 분류 문제로 정의하였다.", indent=True)

add_heading(doc, "6.1 손실 함수", level=2)
add_body(doc,
    "GAT로 학습된 노드 임베딩 h_i, h_j에 대해 내적(dot product)을 이용하여 "
    "링크 존재 확률을 추정하고, Binary Cross-Entropy(BCE) 손실을 최소화한다.")
add_equation(doc,
    "p_ij = sigmoid( h_i^T · h_j )")
add_equation(doc,
    "L_BCE = -[ y_ij · log(p_ij) + (1 - y_ij) · log(1 - p_ij) ]")
add_body(doc,
    "여기서 y_ij = 1이면 positive (실제 엣지), y_ij = 0이면 negative (연결 없음)이다. "
    "약물 노드(210개)의 학습 부족을 방지하기 위해 DGT 엣지(약물-유전자 연결)에 대해 "
    "손실 가중치를 10배로 증가시켰다.")

add_heading(doc, "6.2 Degree 정규화 항", level=2)
add_body(doc,
    "허브 노드의 임베딩이 과도하게 지배적이 되는 것을 방지하기 위해, "
    "연결 차수에 비례하는 L2 정규화 항을 추가하였다.")
add_equation(doc,
    "L = L_BCE + lambda * SUM_v( degree(v) * ||h_v||^2 )")
add_body(doc, "람다(lambda) = 0.01로 설정하였다.")

add_heading(doc, "6.3 학습 설정", level=2)
add_body(doc, "학습에 사용된 주요 하이퍼파라미터는 다음과 같다.")
add_bullet(doc, "옵티마이저: Adam (learning rate = 0.001)")
add_bullet(doc, "최대 에폭: 200 (Early Stopping patience = 20)")
add_bullet(doc, "배치 크기: 4,096 엣지 (양성/음성 1:1 샘플링)")
add_bullet(doc, "Dropout: 0.3 (GAT 레이어 간 적용)")
add_bullet(doc, "가중치 초기화: Xavier Uniform")
add_bullet(doc, "Gradient Clipping: max_norm = 1.0")
add_bullet(doc, "데이터 분할: Train 80% / Validation 10% / Test 10%")

# ════════════════════════════════════════════════
# 7. Attention -> 전이 확률 변환
# ════════════════════════════════════════════════
add_heading(doc, "7. Attention Weight의 RWR 전이 확률 통합", level=1)

add_body(doc,
    "GAT 학습 완료 후, 두 번째 레이어(Layer 2)에서 산출된 방향별 attention 계수 "
    "alpha_ij를 추출하였다. GAT attention은 방향성(i→j vs j→i)을 구분하므로, "
    "각 유향 엣지에 대해 독립적으로 attention 값을 저장하는 directed_attention 딕셔너리를 "
    "구성하였다.", indent=True)

add_heading(doc, "7.1 전이 확률 혼합", level=2)
add_body(doc,
    "GAT attention만을 전이 확률로 직접 사용할 경우, "
    "학습 초기 수렴 불안정에 의해 원본 엣지 가중치 정보가 완전히 손실될 위험이 있다. "
    "이를 방지하기 위해 GAT attention(alpha)과 원본 정규화 가중치(w_orig)를 "
    "α 비율로 선형 혼합하였다.")
add_equation(doc,
    "combined_ij = alpha * GAT_attention_ij  +  (1 - alpha) * w_orig_ij_normalized")
add_equation(doc,
    "P(j | i) = combined_ij / SUM_k( combined_ij )      (local re-normalization)")
add_body(doc,
    "alpha = 0.7로 설정하여 GAT 구조 반영(70%)과 원본 가중치 fallback(30%)을 균형있게 "
    "유지하였다. alpha는 외부 파라미터로 조정 가능하다.")

add_heading(doc, "7.2 이웃 전이 확률 캐시 구성", level=2)
add_body(doc,
    "RWR 시뮬레이션의 효율화를 위해, 모든 노드에 대해 이웃 목록(nbrs)과 "
    "혼합 전이 확률 벡터(combined)를 사전 계산하여 캐시(attention_neighbor_cache)에 저장하였다. "
    "각 워크 스텝에서 확률 배열을 반복 계산하는 비용을 제거함으로써 "
    "RWR 시뮬레이션 속도를 대폭 향상시켰다.", indent=True)

# ════════════════════════════════════════════════
# 8. RWR 시뮬레이션
# ════════════════════════════════════════════════
add_heading(doc, "8. GAT-RWR 시뮬레이션", level=1)

add_body(doc,
    "구성된 attention_neighbor_cache를 이용하여 RWR 시뮬레이션을 수행하였다. "
    "시뮬레이션 구조는 ver1 RWR과 동일하나, 전이 확률로 GAT attention 혼합값을 "
    "사용하는 점이 핵심적인 차이이다.", indent=True)

add_heading(doc, "8.1 이동 규칙", level=2)
add_body(doc, "매 스텝 t에서 보행자는 다음의 규칙을 따른다.")
add_bullet(doc,
    "Teleport (확률 r = 0.15): 시작 노드로 복귀한다. "
    "전역적 탐색 능력을 유지하기 위한 재시작 메커니즘이다.")
add_bullet(doc,
    "이동 (확률 1 - r = 0.85): 전이 확률 캐시의 GAT-혼합 확률에 따라 "
    "이웃 노드 중 하나로 이동한다.")

add_heading(doc, "8.2 시뮬레이션 설정", level=2)
add_bullet(doc, "시작 노드: 전체 노드 대상 (|V| = 18,211개)")
add_bullet(doc, "워크 횟수(n_walks): 시작 노드당 1,000회")
add_bullet(doc, "워크 길이(walk_length): 워크당 최대 100 스텝")
add_bullet(doc, "Restart 확률(r): 0.15")
add_bullet(doc, "병렬 워커 수: 12 (Python multiprocessing.Pool)")
add_bullet(doc, "난수 시드: 42 (재현성 확보)")

add_heading(doc, "8.3 노드 중요도 정의", level=2)
add_body(doc,
    "GAT-RWR 시뮬레이션 결과 각 노드의 방문 횟수(visit_count)를 집계하고, "
    "전체 방문 횟수의 합으로 나누어 방문 확률(visit_prob)을 산출한다.")
add_equation(doc, "Importance_GAT-RWR(v) = visit_count(v) / SUM_u( visit_count(u) )")
add_body(doc,
    "방문 확률이 높은 노드는 GAT가 학습한 네트워크 구조적 관점에서도 "
    "다수의 이웃과 강한 attention 연결을 가지며, 여러 경로에서 빈번히 경유되는 "
    "구조적 허브이다.")

# ════════════════════════════════════════════════
# 9. GAT Hub Score
# ════════════════════════════════════════════════
add_heading(doc, "9. GAT Hub Score 산출", level=1)

add_body(doc,
    "노드 v에 대한 GAT Hub Score는 해당 노드를 목적지(destination)로 하는 "
    "모든 incoming attention 계수의 평균으로 정의한다.")
add_equation(doc,
    "gat_hub_score(v) = (1 / |IN(v)|) * SUM_{u in IN(v)}( alpha_uv )")
add_body(doc,
    "여기서 IN(v)는 v로 incoming 엣지를 갖는 이웃 집합이다. "
    "gat_hub_score는 단순 방문 확률 기반 중요도(visit_prob)와 상보적인 정보를 제공하며, "
    "다른 노드들이 v를 얼마나 중요하게 '주목(attend)'하는지를 나타낸다.", indent=True)

# ════════════════════════════════════════════════
# 10. 결과 저장 및 분석
# ════════════════════════════════════════════════
add_heading(doc, "10. 결과 저장 및 분석", level=1)

add_body(doc,
    "GAT-RWR 시뮬레이션 결과는 node_importance_v2.tsv 파일에 저장된다. "
    "출력 파일은 다음의 컬럼을 포함한다.", indent=True)
add_bullet(doc, "rank: 방문 확률 기준 전체 순위")
add_bullet(doc, "node: 노드명 (유전자 심볼 또는 약물명)")
add_bullet(doc, "node_type: gene / drug")
add_bullet(doc, "visit_count: RWR 시뮬레이션 총 방문 횟수")
add_bullet(doc, "visit_prob: 방문 확률 (전체 방문 중 해당 노드 비율)")
add_bullet(doc, "degree: 연결 차수")
add_bullet(doc, "gat_hub_score: 해당 노드로 향하는 incoming attention 평균")

add_body(doc,
    "ver1 RWR 결과(node_importance.tsv)와 ver2 GAT-RWR 결과(node_importance_v2.tsv)를 "
    "비교함으로써, GAT attention 통합이 핵심 노드 순위에 미치는 영향을 분석하였다. "
    "특히 약물 노드의 순위 변화는 GAT가 약물-유전자 표적 관계의 네트워크 맥락을 "
    "추가적으로 반영함으로써 치료 후보 발굴에 기여하는 정도를 나타낸다.", indent=True)

# ════════════════════════════════════════════════
# 11. 구현 세부사항
# ════════════════════════════════════════════════
add_heading(doc, "11. 구현 세부사항 및 수치 안정화", level=1)

add_body(doc,
    "대규모 그래프(18,211 노드, 288,336 엣지) 환경에서의 수치 안정성과 "
    "연산 효율을 위해 다음의 기법들을 적용하였다.", indent=True)

add_bullet(doc,
    "Softmax 수치 안정화: e_ij에서 max(e_ij)를 차감한 후 softmax를 적용하여 "
    "overflow 방지 (numerically stable softmax)")
add_bullet(doc,
    "Attention 하한값: alpha < 1e-8인 경우 1e-8로 클리핑하여 "
    "zero-division 방지")
add_bullet(doc,
    "scatter_add 구현: PyG 없이 PyTorch의 Tensor.scatter_add_()로 "
    "이웃 aggregation 직접 구현")
add_bullet(doc,
    "멀티프로세싱: multiprocessing.Pool + initializer 패턴을 이용하여 "
    "공유 진행률 카운터를 워커에 전달 (Windows pickle 문제 해결)")
add_bullet(doc,
    "Progress bar: ASCII 문자(#/-)를 사용하여 Windows cp949 인코딩 환경에서의 "
    "Unicode 오류 방지")

# ════════════════════════════════════════════════
# 저장
# ════════════════════════════════════════════════
out_path = "C:/Users/user/Desktop/CNN_CVPR/__DATAAn__/random_work_test/method_GAT-RWR.docx"
doc.save(out_path)
print(f"[Done] 저장 완료: {out_path}")
